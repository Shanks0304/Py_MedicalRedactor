from datetime import datetime
import json
import os
from pathlib import Path
import platform
import re
import sys
from docx import Document
from PyQt6.QtCore import QObject, pyqtSignal
from dotenv import load_dotenv
from pydantic import BaseModel
import multiprocessing

from utils.config import setup_logger

load_dotenv()

class LLMService(QObject):
    response_ready = pyqtSignal(str)
    debug_message = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self):
        super().__init__()
        
        self.logger = setup_logger(__name__)
        
        if getattr(sys, 'frozen', False):
            # If running from bundle
            if platform.system() == 'Darwin':  # macOS
                bundle_dir = os.path.dirname(sys.executable)
                resources_dir = os.path.join(os.path.dirname(bundle_dir), 'Resources')
                dotenv_path = os.path.join(resources_dir, '.env')
            else:  # Windows
                bundle_dir = os.path.dirname(sys.executable)
                dotenv_path = os.path.join(bundle_dir, '_internal', '.env')
                
            self.logger.info(f"Running as bundled app on {platform.system()}")
            self.logger.info(f"Loading .env from: {dotenv_path}")
        else:
            # If running from source
            dotenv_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), '.env')
            
        # Load environment variables
        if os.path.exists(dotenv_path):
            load_dotenv(dotenv_path)
        else:
            raise Exception(f".env file not found at {dotenv_path}")
        

        self.logger.info(f"Loading .env from: {dotenv_path}")
        self.logger.info(f"Exists: {os.path.exists(dotenv_path)}")
        self.logger.info(f"OPENAI_API_KEY set: {'OPENAI_API_KEY' in os.environ}")

        self.output_dir = os.path.join(os.path.expanduser('~/Documents'), 'medicalapp', 'llm_outputs')
        os.makedirs(self.output_dir, exist_ok=True)
        
        self.template_doc = None
        self.template_structure = None
        self.current_response_doc = None
        self.llm = None

        
        load_dotenv(dotenv_path)

    def set_template(self, template_path: str):
        """Load and analyze the template document"""
        try:
            if not Path(template_path).exists():
                raise FileNotFoundError(f"Template file not found: {template_path}")
                
            self.template_doc = Document(template_path)
            
            # Validate template
            if not self.template_doc.paragraphs:
                raise ValueError("Template document is empty")
                
            return True
            
        except Exception as e:
            error_msg = f"Error loading template: {str(e)}"
            self.logger.error(error_msg)  # Debug print
            self.error_occurred.emit(error_msg)
            return False

    def process_text(self, transcription: str):
        """Process the transcription with the loaded template"""
        try:
            if not transcription:
                raise ValueError("No transcription provided")
            
            if not self.template_doc:
                raise ValueError("No template loaded")

            from presidio_analyzer import AnalyzerEngine
            analyzer = AnalyzerEngine()
            analyzer_results = analyzer.analyze(text=transcription, language='en')

            entity_counters = {}
            replacements = []
            patient_data = []

            self.logger.info("Identified these PII entities:")
            for result in analyzer_results:
                if result.score > 0.5:
                    # Initialize counter for new entity types
                    if result.entity_type not in entity_counters:
                        entity_counters[result.entity_type] = 1
                    
                    # Create replacement with indexed entity type
                    replacement = {
                        'start': result.start,
                        'end': result.end,
                        'original': transcription[result.start:result.end],
                        'replacement': f"{{{result.entity_type}_{entity_counters[result.entity_type]}}}"
                    }
                    replacements.append(replacement)
                    
                    # Increment counter for this entity type
                    entity_counters[result.entity_type] += 1

                    self.logger.info(f"- {transcription[result.start:result.end]} as {result.entity_type}")
                    patient_data.append({f"{{{result.entity_type}_{entity_counters[result.entity_type]}}}": transcription[result.start:result.end]})

            self.logger.info(f"Patient data: {patient_data}")
            # Apply replacements from end to start to avoid index shifting
            updated_transcription = transcription
            for replacement in sorted(replacements, key=lambda x: x['start'], reverse=True):
                updated_transcription = (
                    updated_transcription[:replacement['start']] + 
                    replacement['replacement'] + 
                    updated_transcription[replacement['end']:]
                )

            print("Modified text:")
            print(updated_transcription)

            sample_note = '\n'.join([paragraph.text for paragraph in self.template_doc.paragraphs])

            # Get list of placeholder keys from patient_data
            placeholders = []
            for data_dict in patient_data:
                placeholders.extend(list(data_dict.keys()))

            import openai
            response = openai.chat.completions.create(
                model="gpt-4o-mini",
            #     # messages=[
            #     #     {"role": "system", "content": "You are a bot that rewrites medical consultation notes using provided interview transcripts, ensuring the use of {name}, {age}, {date_of_birth}, {email}, {phone_number}."},
            #     #     {"role": "user", "content":
            #     #         f"""
            #     #         Here is the original consultation note template:
            #     #         {sample_note}

            #     #         Please rewrite this consultation note using the following interview transcription:
            #     #         {updated_transcription}"""
            #     #         + """
            #     #         Important:

            #     #         Find all the available and possible patient's name, age, email, phone number and date of birth.

            #     #         Replace them with {name}, {age}, {date_of_birth}, {email}, {phone_number} on new note.

            #     #         Don't use any additional or creative labels except for {name}, {age}, {date_of_birth}, {email}, {phone_number}.

            #     #         Do not retain any previous PII from the sample note.

            #     #         Retain the original structure of the consultation note and do not include any explanations, comments, or additional text.
            #     #         """
            #     #     }
            #     # ]
                messages = [
                    {
                        "role": "system", "content":
                            """You are a professional medical documentation assistant specializing in rewriting consultation notes.
                            Your role is to synthesize clinical interview transcripts into polished, human-like consultation notes that are comprehensive, detailed, and empathetic while adhering to a professional structure.
                            Ensure that sensitive information like HIPAA compliance identifying data are replaced with brackets like {name}, {age}, {date_of_birth}.
                            Your notes should accurately reflect the patient’s symptoms, history, and context in a nuanced and relatable manner.
                            Only return the same structured rewritten consultation notes without any additional information, style and formatting.
                            """
                    },
                    {
                        "role": "user",
                        "content": f"""
                        Here is the original consultation note template:
                        {sample_note}

                        Please rewrite this consultation note using the following interview transcription:
                        {updated_transcription}

                        These are placeholders that should be used on generated document: {placeholders}
                        Use only these placeholders.
                        
                        Only return the same structured rewritten consultation notes without any additional information, style and formatting.
                        """ + 

                        
                        """
                        Important:
                        1. Retain the structure of the provided consultation note template, including all sections and subheadings.
                        2. Ensure that all details mentioned in the transcription (e.g., demographic information, symptoms, history) are incorporated accurately into the rewritten note. Avoid omitting any relevant information unless it is not present in the transcription.
                        3. Describe the patient’s symptoms in a human-like and relatable manner. For example:
                        - Instead of stating "The patient reports a long-standing low mood," elaborate with, "The patient shared that their mood has felt persistently low, describing it as a constant ‘cloud’ that dampens their day-to-day experiences."
                        - Include the patient's own words, where appropriate, to make the note more vivid.
                        4. Provide clear and professional clinical reasoning in sections like ‘Impression’ and ‘Plan,’ integrating the patient’s history and symptoms into the assessment.
                        5. Avoid over-reliance on placeholders (e.g., {name}) when the same placeholder is repeatedly used within the same section. Use pronouns appropriately for natural readability.
                        6. Do not simply copy and paste verbatim from the provided template or transcript. Rewrite to ensure the note feels cohesive and thoughtfully composed.
                        7. Avoid clinical jargon unless absolutely necessary, opting for plain, professional language that is accessible and clear.



                        The rewritten note should present as a polished, nuanced, and complete consultation note, avoiding redundancy or omissions. It should read as though written by a highly experienced and empathetic clinician."""
                    }
                ]
            )
            
            new_note = response.choices[0].message.content
            self.logger.info("\n" + new_note)

            new_note = self.replace_pii_with_labels(new_note, patient_data)

            self.response_ready.emit(new_note)
            return new_note
                    
        except Exception as e:
            self.error_occurred.emit(str(e))
            return None

    def save_response(self, response_context: str) -> str:
        """
        Save the response to a Word document in the results directory at project root
        
        Args:
            response_context (str): The content to save
            
        Returns:
            str: Path to saved file, or None if error occurs
        """
        try:
            if not response_context:
                raise ValueError("No content to save")

            # Get project root directory and create results folder
            results_dir = os.path.join(os.path.expanduser('~/Documents'), 'medicalapp', 'results') # Go up to project root
            os.makedirs(results_dir, exist_ok=True)

            # Create new document and add content
            doc = Document()
            doc.add_paragraph(response_context)

            # Create filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"consultation_note_{timestamp}.docx"
            output_path = results_dir / filename

            # Save the document
            doc.save(str(output_path))
            self.logger.info(f"Document saved successfully at: {output_path}")
            return str(output_path)

        except Exception as e:
            error_msg = f"Error saving response: {str(e)}"
            self.logger.error(error_msg)  # Debug print
            self.error_occurred.emit(error_msg)
            return None
    
    def replace_pii_with_labels(self, text: str, replacements: list) -> str:
        # Merge all dictionaries in the list into a single dictionary
        merged_replacements = {}
        for d in replacements:
            merged_replacements.update(d)
        
        
        # Convert all keys to lowercase and values to strings
        merged_replacements = {key: str(value) for key, value in merged_replacements.items()}
        self.logger.info(f"Merged replacements: {merged_replacements}")
        
        # This regex finds text in curly braces
        pattern = r'\{[A-Z_0-9]+\}'
        
        # Replace each occurrence with its corresponding value
        result = text
        for key, value in merged_replacements.items():
            result = result.replace(key, value)
        
        return result